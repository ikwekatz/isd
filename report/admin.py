from io import BytesIO

from django.contrib import admin
from django import forms
from django.db.models import Sum
from django.urls import path
from django.http import HttpResponse
from django.template.loader import render_to_string
import datetime
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.http import HttpResponse

from activities.models import FinancialYear, Activity, Budget, Expenditure
from office.models import Unit, Section
from services.models import SupportService, SupportTicket, StatisticType, StatisticsRecord
from django.db import models
class Report(Activity):
    class Meta:
        proxy = True
        verbose_name = 'Activity Report'
        verbose_name_plural = 'Activity Reports'

class ReportForm(forms.Form):
    UNIT_SECTION_CHOICES = [
        ('unit', 'By Unit'),
        ('section', 'By Section'),
    ]

    grouping = forms.ChoiceField(choices=UNIT_SECTION_CHOICES, required=True)
    unit = forms.ModelChoiceField(queryset=Unit.objects.all(), required=False)
    section = forms.ModelChoiceField(queryset=Section.objects.all(), required=False)  # Fixed: changed models to forms
    start_date = forms.DateField(required=True, widget=forms.DateInput(attrs={'type': 'date'}))
    end_date = forms.DateField(required=True, widget=forms.DateInput(attrs={'type': 'date'}))
    financial_year = forms.ModelChoiceField(queryset=FinancialYear.objects.all(), required=True)

    def clean(self):
        cleaned_data = super().clean()
        grouping = cleaned_data.get('grouping')
        unit = cleaned_data.get('unit')
        section = cleaned_data.get('section')

        if grouping == 'unit' and not unit:
            raise forms.ValidationError("Please select a unit when grouping by unit")
        if grouping == 'section' and not section:
            raise forms.ValidationError("Please select a section when grouping by section")

        start_date = cleaned_data.get('start_date')
        end_date = cleaned_data.get('end_date')

        if start_date and end_date and end_date < start_date:
            raise forms.ValidationError("End date cannot be earlier than start date")

        return cleaned_data

def set_cell_width(cell, width):  # width must be Inches object or float inches
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing width tags
    for child in tcPr.findall(qn('w:tcW')):
        tcPr.remove(child)
    w = OxmlElement('w:tcW')
    w.set(qn('w:type'), 'dxa')

    # If width is Inches, convert to dxa; if float assume inches
    if hasattr(width, 'inches'):
        width_value = int(width.inches * 1440)
    else:
        width_value = int(width * 1440)  # width is float inches

    w.set(qn('w:w'), str(width_value))
    tcPr.append(w)

class ReportAdmin(admin.ModelAdmin):
    change_list_template = 'admin/report_change_list.html'

    def has_add_permission(self, request):
        return False

    def has_change_permission(self, request, obj=None):
        return False

    def has_delete_permission(self, request, obj=None):
        return False

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('generate-report/', self.admin_site.admin_view(self.generate_report), name='generate_report'),
            path('export-word/', self.admin_site.admin_view(self.export_word), name='export_word'),

        ]
        return custom_urls + urls

    def changelist_view(self, request, extra_context=None):
        extra_context = extra_context or {}
        extra_context['report_form'] = ReportForm()
        return super().changelist_view(request, extra_context=extra_context)

    def generate_report(self, request):
        if request.method != 'POST':
            return HttpResponse("Method not allowed", status=405)

        form = ReportForm(request.POST)
        if not form.is_valid():
            return HttpResponse("Invalid form data", status=400)

        data = form.cleaned_data
        grouping = data['grouping']
        unit = data.get('unit')
        section = data.get('section')
        start_date = data['start_date']
        end_date = data['end_date']
        financial_year = data['financial_year']

        end_date = end_date + datetime.timedelta(days=1)

        if grouping == 'unit':
            activities = Activity.objects.filter(unit=unit, financial_year=financial_year)
        else:
            activities = Activity.objects.filter(section=section, financial_year=financial_year)

        report_data = []

        for activity in activities:
            # Implementation part (unchanged)
            support_services = SupportService.objects.filter(activities=activity)
            ticket_implementations = []
            for service in support_services:
                tickets = SupportTicket.objects.filter(service=service,
                                                       submitted_at__gte=start_date,
                                                       submitted_at__lt=end_date)
                if not tickets.exists():
                    continue
                ticket_implementations.append({
                    'service': service.name,
                    'description': "\n".join(f"- {t.description} (Status: {t.get_status_display()})" for t in tickets)
                })

            # Statistics fallback
            if not ticket_implementations:
                statistic_types = StatisticType.objects.filter(activities=activity)
                stats = StatisticsRecord.objects.filter(statistic_type__in=statistic_types,
                                                        start_date__gte=start_date,
                                                        end_date__lte=end_date)
                if stats.exists():
                    description = "\n".join(f"- {stat.title}: {stat.description}" for stat in stats)
                    ticket_implementations.append({
                        'service': "Statistics Records",
                        'description': description
                    })

            # --- Financial data ---
            budgets_qs = Budget.objects.filter(activity=activity, financial_year=financial_year) \
                .values('budget_type') \
                .annotate(total_budget=Sum('amount'))
            expenditures_qs = Expenditure.objects.filter(activity=activity, financial_year=financial_year) \
                .values('budget_type') \
                .annotate(total_expenditure=Sum('amount'))

            # Prepare dictionaries per budget type
            budgets = list(budgets_qs)
            exp_map = {e['budget_type']: e['total_expenditure'] for e in expenditures_qs}
            bal_map = {b['budget_type']: b['total_budget'] - exp_map.get(b['budget_type'], 0) for b in budgets}
            total_budget = sum(b['total_budget'] for b in budgets)
            total_expenditure = sum(exp_map.values())
            total_balance = total_budget - total_expenditure

            report_data.append({
                'activity': activity,
                'implementations': ticket_implementations,
                'budgets': budgets,
                'exp_map': exp_map,
                'balance_map': bal_map,
                'total_budget': total_budget,
                'total_expenditure': total_expenditure,
                'balance_total': total_balance,
            })

        context = {
            'report_data': report_data,
            'grouping': grouping,
            'unit': unit,
            'section': section,
            'start_date': start_date,
            'end_date': end_date - datetime.timedelta(days=1),
            'financial_year': financial_year,
            'generated_on': datetime.datetime.now(),
            'total_budget': sum(item['total_budget'] for item in report_data),
            'total_expenditure': sum(item['total_expenditure'] for item in report_data),
            'total_balance': sum(item['balance_total'] for item in report_data),
        }

        return HttpResponse(render_to_string('admin/report_output.html', context))


    def export_word(self, request):
        if request.method != 'POST':
            return HttpResponse("Method not allowed", status=405)

        form = ReportForm(request.POST)
        if not form.is_valid():
            return HttpResponse("Invalid form data", status=400)

        data = form.cleaned_data
        grouping = data['grouping']
        unit = data.get('unit')
        section = data.get('section')
        start_date = data['start_date']
        end_date = data['end_date']
        financial_year = data['financial_year']

        # Include full last day
        end_date = end_date + datetime.timedelta(days=1)

        # Get activities based on grouping
        if grouping == 'unit':
            activities = Activity.objects.filter(unit=unit, financial_year=financial_year)
        else:
            activities = Activity.objects.filter(section=section, financial_year=financial_year)

        doc = Document()
        sections = doc.sections[0]
        header = sections.header
        header_para = header.paragraphs[0]

        if grouping == 'unit':
            unit_name = unit.name if unit else 'N/A'
            heading = f"\t {unit_name} Activities Implementation Report from {start_date} to {end_date - datetime.timedelta(days=1)} in Financial Year: {financial_year}"
            header_para.text = heading
        else:
            section_name = section.name if section else 'N/A'
            doc.add_heading(f"{section_name} Activities Implementation Report {start_date} to {end_date - datetime.timedelta(days=1)} in Financial Year: {financial_year}", level=0)

        # Create table with extra columns
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'SN'
        hdr_cells[1].text = 'Activity'
        hdr_cells[2].text = 'Implementation'
        hdr_cells[3].text = 'Budget'
        hdr_cells[4].text = 'Expenditure'
        hdr_cells[5].text = 'Balance'

        total_width_inches = 9
        col_widths = [
            Inches(total_width_inches * 0.05),  # SN
            Inches(total_width_inches * 0.20),  # Activity
            Inches(total_width_inches * 0.35),  # Implementation
            Inches(total_width_inches * 0.13),  # Budget
            Inches(total_width_inches * 0.13),  # Expenditure
            Inches(total_width_inches * 0.14),  # Balance
        ]
        for cell, width in zip(hdr_cells, col_widths):
            set_cell_width(cell, width)

        for idx, activity in enumerate(activities, 1):
            support_services = SupportService.objects.filter(activities=activity)
            ticket_implementations = []
            for service in support_services:
                tickets = SupportTicket.objects.filter(
                    service=service,
                    submitted_at__gte=start_date,
                    submitted_at__lt=end_date
                )
                if not tickets.exists():
                    continue
                ticket_implementations.append({
                    'service': service.name,
                    'description': "\n".join(f"- {t.description} (Status: {t.get_status_display()})" for t in tickets)
                })

            # ---- Financial Data ----
            budgets = Budget.objects.filter(activity=activity, financial_year=financial_year) \
                .values("budget_type") \
                .annotate(total_budget=Sum("amount"))

            total_budget = sum(b["total_budget"] for b in budgets)

            expenditures = Expenditure.objects.filter(activity=activity, financial_year=financial_year) \
                .values("budget_type") \
                .annotate(total_expenditure=Sum("amount"))

            exp_map = {e["budget_type"]: e["total_expenditure"] or 0 for e in expenditures}
            total_expenditure = sum(exp_map.values())

            balance_map = {b["budget_type"]: b["total_budget"] - exp_map.get(b["budget_type"], 0) for b in budgets}
            balance_total = total_budget - total_expenditure

            # Add row
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(activity)

            # Implementation cell
            impl_cell = row_cells[2]
            if ticket_implementations:
                for impl in ticket_implementations:
                    impl_cell.add_paragraph(impl['service'], style='Heading4')
                    impl_cell.add_paragraph(impl['description'])
            else:
                impl_cell.text = "No records"

            # Budget
            budget_cell = row_cells[3]
            for b in budgets:
                budget_cell.add_paragraph(f"{b['budget_type']}: {b['total_budget']}")
            budget_cell.add_paragraph(f"TOTAL: {total_budget}")

            # Expenditure
            exp_cell = row_cells[4]
            for b in budgets:
                exp_cell.add_paragraph(f"{b['budget_type']}: {exp_map.get(b['budget_type'], 0)}")
            exp_cell.add_paragraph(f"TOTAL: {total_expenditure}")

            # Balance
            bal_cell = row_cells[5]
            for b in budgets:
                bal_cell.add_paragraph(f"{b['budget_type']}: {balance_map.get(b['budget_type'], 0)}")
            bal_cell.add_paragraph(f"TOTAL: {balance_total}")

            # Adjust column widths
            for cell, width in zip(row_cells, col_widths):
                set_cell_width(cell, width)

        # Export DOCX
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        buffer.close()
        return response


    # def export_word(self, request):
    #     if request.method != 'POST':
    #         return HttpResponse("Method not allowed", status=405)
    #
    #     form = ReportForm(request.POST)
    #     if not form.is_valid():
    #         return HttpResponse("Invalid form data", status=400)
    #
    #     data = form.cleaned_data
    #     grouping = data['grouping']
    #     unit = data.get('unit')
    #     section = data.get('section')
    #     start_date = data['start_date']
    #     end_date = data['end_date']
    #     financial_year = data['financial_year']
    #
    #     # Include the full last day
    #     end_date = end_date + datetime.timedelta(days=1)
    #
    #     # Get activities based on grouping
    #     if grouping == 'unit':
    #         activities = Activity.objects.filter(
    #             unit=unit,
    #             financial_year=financial_year
    #         )
    #     else:
    #         activities = Activity.objects.filter(
    #             section=section,
    #             financial_year=financial_year
    #         )
    #
    #     doc = Document()
    #     sections = doc.sections[0]
    #     header = sections.header
    #     header_para = header.paragraphs[0]
    #
    #     if grouping == 'unit':
    #         unit_name = unit.name if unit else 'N/A'
    #         heading = f"\t {unit_name} Activities Implementation Report from {start_date} to {end_date - datetime.timedelta(days=1)} in Financial Year: {financial_year} "
    #         header_para.text = heading
    #     else:
    #         section_name = section.name if section else 'N/A'
    #         doc.add_heading(f"{section_name} Activities Implementation Report {start_date} to {end_date - datetime.timedelta(days=1)} in Financial Year: {financial_year}", level=0)
    #
    #     table = doc.add_table(rows=1, cols=3)
    #     table.style = 'Table Grid'
    #     hdr_cells = table.rows[0].cells
    #     hdr_cells[0].text = 'SN'
    #     hdr_cells[1].text = 'Activity'
    #     hdr_cells[2].text = 'Implementation'
    #     total_width_inches = 8
    #     col_widths = [
    #         Inches(total_width_inches * 0.05),
    #         Inches(total_width_inches * 0.25),
    #         Inches(total_width_inches * 0.70),
    #     ]
    #     for cell, width in zip(hdr_cells, col_widths):
    #         set_cell_width(cell, width)
    #
    #     report_data = []
    #     for activity in activities:
    #         support_services = SupportService.objects.filter(activities=activity)
    #         ticket_implementations = []
    #
    #         for service in support_services:
    #             tickets = SupportTicket.objects.filter(
    #                 service=service,
    #                 submitted_at__gte=start_date,
    #                 submitted_at__lt=end_date
    #             ).order_by('service__name')
    #
    #             if not tickets.exists():
    #                 continue
    #             ticket_count = tickets.count()
    #             service_display_name = f"{service.name} (Occurrence: {ticket_count})"
    #             if service.is_related_to_system:
    #                 systems = set(t.system for t in tickets if t.system)
    #                 for system in systems:
    #                     system_tickets = [t for t in tickets if t.system == system]
    #                     description = "".join(
    #                         f"- {t.description}  User: {t.internal_user_name or ''} {t.external_user or ''}  (Status: {t.get_status_display()})"
    #                         for t in system_tickets
    #                     )
    #                     ticket_implementations.append({
    #                         'service': service_display_name,
    #                         'description': f"System: {system}\n{description}" if system else description
    #                     })
    #             else:
    #                 description = "".join(
    #                     f"- {t.description} (Status: {t.get_status_display()})"
    #                     for t in tickets
    #                 )
    #                 ticket_implementations.append({
    #                     'service': service_display_name,
    #                     'description': description
    #                 })
    #
    #         if not ticket_implementations:
    #             statistic_types = StatisticType.objects.filter(activities=activity)
    #             stats = StatisticsRecord.objects.filter(
    #                 statistic_type__in=statistic_types,
    #                 start_date__gte=start_date,
    #                 end_date__lte=end_date
    #             )
    #
    #             if stats.exists():
    #                 description = "".join(
    #                     f"- {stat.title}: {stat.description}"
    #                     for stat in stats
    #                 )
    #                 ticket_implementations.append({
    #                     'service': "Statistics Records",
    #                     'description': description
    #                 })
    #
    #         if ticket_implementations:
    #             report_data.append({
    #                 'activity': activity,
    #                 'implementations': ticket_implementations
    #             })
    #
    #     for idx, item in enumerate(report_data, 1):
    #         row_cells = table.add_row().cells
    #         row_cells[0].text = str(idx)
    #         row_cells[1].text = str(item['activity'])
    #         for cell, width in zip(row_cells, col_widths):
    #             set_cell_width(cell, width)
    #
    #         impl_cell = row_cells[2]
    #         for impl in item['implementations']:
    #             impl_cell.add_paragraph(impl['service'], style='Heading4')
    #             impl_cell.add_paragraph(impl['description'])
    #     buffer = BytesIO()
    #     doc.save(buffer)
    #     buffer.seek(0)
    #     filename = f"unit_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    #     response = HttpResponse(
    #         buffer.getvalue(),
    #         content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    #     )
    #     response['Content-Disposition'] = f'attachment; filename="{filename}"'
    #     buffer.close()
    #     return response

@admin.register(Report)
class ReportAdminView(ReportAdmin):
    pass